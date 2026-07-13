"""docx_extractor: DOCX -> clean text (python-docx).

Isolation-first tool (see tools/README.md). Allowed: python-docx + stdlib
+ optional structured_logger. No orchestrator/sleep_*/redis_*/engine/mtr_*.
"""
from __future__ import annotations
import os, sys

try:
    from structured_logger import get_event_logger
    _logger = get_event_logger("docx_extractor")
except Exception:
    _logger = None

_EXTS = (".docx",)


def convert_docx_to_markdown(input_path: str, output_path: str) -> None:
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"input file not found: {input_path}")
    if not input_path.lower().endswith(_EXTS):
        raise ValueError(f"input is not a .docx file: {input_path}")

    if _logger:
        _logger.log(event_type="extraction_started", data={"source": input_path})

    try:
        from docx import Document
        doc = Document(input_path)
    except Exception as e:
        raise RuntimeError(f"DOCX parse failed: {e}") from e

    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
    text = _normalize("\n".join(parts))  # empty doc allowed
    _write(text, input_path, output_path)


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.rstrip() for ln in text.split("\n")]
    out = "\n".join(lines)
    while "\n\n\n" in out:
        out = out.replace("\n\n\n", "\n\n")
    return out.strip() + "\n" if out.strip() else ""


def _write(text: str, input_path: str, output_path: str) -> None:
    out_dir = os.path.dirname(os.path.abspath(output_path))
    try:
        os.makedirs(out_dir, exist_ok=True)
    except OSError as e:
        raise IOError(f"cannot create output directory {out_dir}: {e}") from e
    try:
        with open(output_path, "w", encoding="utf-8") as fh:
            fh.write(text)
    except OSError as e:
        raise IOError(f"cannot write output {output_path}: {e}") from e
    if _logger:
        _logger.log(event_type="extraction_complete",
                    data={"source": input_path, "dest": output_path, "char_count": len(text)})
    print(f"Converted {input_path} → {output_path} ({len(text)} chars)", file=sys.stdout)
